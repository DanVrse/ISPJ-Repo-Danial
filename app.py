
import pyotp
import qrcode
import psutil
# from Image import *


from watermarking import startwatermark,decodewatermark
from flask import Flask, url_for,abort, render_template, request, redirect, jsonify, send_from_directory, \
    current_app, make_response, session, send_file
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer

import random

# file upload
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import os

import Encryption as E
# import magic
import urllib.request
from datetime import datetime
import time
import json
import re
# Credentials
from werkzeug.security import generate_password_hash, check_password_hash
from fileSystemOrganisation import *

from random import randint
from SQLite_Functions import *
from User import OrgUser, LogUser
from db_restore import *
from orgBackup import *
import threading
import socket
import sqlite3
import mailingkc as mkc
import io
from docx import Document  # pip install python-docx
import shutil

app = Flask(__name__)

UPLOAD_FOLDER = '/static'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['UPLOAD_EXTENSIONS'] = ['.jpg', '.png']
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024


UPLOAD_FOLDER = './ISPJ uploads'
BASE_DIR = './ISPJ uploads/user'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

app.secret_key = 'Infosecurity_Project'
app.config['ENV'] = 'development'
app.config['DEBUG'] = True
app.config['TESTING'] = True

# MAIL
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_DEBUG'] = True
app.config['MAIL_USERNAME'] = 'ShakeShackproject@gmail.com'
app.config['MAIL_PASSWORD'] = 'Shakeshackproject123'
app.config['MAIL_DEFAULT_SENDER'] = None
app.config['MAIL_MAX_EMAILS'] = None
app.config['MAIL_SUPPRESS_SEND'] = False
app.config['MAIL_ASCII_ATTACHMENTS'] = False
ENABLE_MAIL = True
mail = Mail(app)

# Change this to True, to automate backup every app launch.
VERIFY_FILES = False


app.secret_key = 'INFOSECURITY_PROJECT'
app.config['SECURITY_PASSWORD_SALT'] = 'peace_sign'
app.config['ENV'] = 'development'
app.config['DEBUG'] = True
app.config['TESTING'] = True


sessionDict = {}

uploadpath = None
url = None
seperator = "!$%()"
permseperator = "<|>"
userole = None
mailfea = "yes"
certdict = {}

# Creating token
def generate_confirmation_token(email):
    serializer = URLSafeTimedSerializer(app.secret_key)
    return serializer.dumps(email, salt=app.config['SECURITY_PASSWORD_SALT'])


# Confirms the token that is created.
def confirm_token(token, expiration=3600):
    serializer = URLSafeTimedSerializer(app.secret_key)
    try:
        email = serializer.loads(
            token,
            salt=app.config['SECURITY_PASSWORD_SALT'],
            max_age=expiration
        )

    except:
        return False
    return email


def password_policy(password):
    """
    Verify the strength of 'password'
    Returns a dict indicating the wrong criteria
    A password is considered strong if:
        8 characters length or more
        1 digit or more
        1 symbol or more
        1 uppercase letter or more
        1 lowercase letter or more
    """

    # calculating the length
    short_length_error = len(password) < 8

    long_length_error = len(password) > 256
    # searching for digits
    digit_error = re.search(r"\d", password) is None

    # searching for uppercase
    uppercase_error = re.search(r"[A-Z]", password) is None

    # searching for lowercase
    lowercase_error = re.search(r"[a-z]", password) is None

    # searching for symbols

    acceptable = list(map(chr, range(48, 58))) + list(map(chr, range(65, 91))) + list(map(chr, range(97, 123))) \
                 + list(map(chr, (33, 35, 36, 64, 95, 126)))
    match = [characters in acceptable for characters in password]
    if False in match or len(match) == 0:
        symbol_error = True
    else:
        symbol_error = False

    # overall result
    password_ok = not (
            short_length_error or long_length_error or digit_error or uppercase_error or lowercase_error or symbol_error)

    return {
        'password_ok': password_ok,
        'Short_length_error': short_length_error,
        'Long_length_error': long_length_error,
        'digit_error': digit_error,
        'uppercase_error': uppercase_error,
        'lowercase_error': lowercase_error,
        'symbol_error': symbol_error,
    }


# Generates a session ID after Login + 2FA
def generateSessionID():
    sessionID = ''
    for _ in range(10):
        sessionID += str(random.randint(1, 9))
    return sessionID


@app.before_request
def before_req():
    # session.clear()
    global userrole
    print(session)
    if request.endpoint not in ["login"]:
        session["reqorg"] =None
    try:

        print(session["root"],"YAYYY")
        if session["root"] != None:
            userrole = request.cookies.get('role')
            print(userrole,"LLLLLLLLLLLLL")
            if request.endpoint not in ["home", "createfolder","upload","permission"]:
                session['filepage'] = f"{BASE_DIR}/{session['root']}"


            if len(session["filepage"]) < len(f"{BASE_DIR}/{session['root']}"):
                session["filepage"] = f"./ISPJ uploads/user/{session['root']}"
    except Exception as e:
        session["root"]=None
        print(e)


@app.route('/', methods=['GET', 'POST'])
def login():
    global certdict
    if 'sessionID' in session:
        return redirect(url_for('home'))
    session.clear()
    session["root"]=None
    try:
        reqorg = session["reqorg"]
    except Exception:
        reqorg = None
    if request.method == 'POST':
        print('/ POST RECEIVED')
        email = request.form['email']
        password = request.form['password']
        cert = request.files['fileUpload']
        # with open('users.json', 'r') as f:
        #     usersDict = json.load(f)

        # Extracting organization folders.
        folders = os.listdir(BASE_DIR)

        organisation = request.form['organization']
        organisation = seperator.join(organisation.split("#"))

        print(organisation,"<- ORGANISATION")

        # If the organization field is EMPTY.
        if organisation == "":
            for element in folders:
                print(element, email)
                if email in element:
                    organisation = element

                    orgPath = os.path.join(BASE_DIR, element)
                    existEmail = db_Login(email, orgPath)
                    emailCheck = db_Query(email, orgPath, column='email')
                    # db_Query(email, f"{BASE_DIR}/{organisation}/", column='email')
                    passCheck = db_Query(email, orgPath, column='password')
                    break
                else:
                    print("ERROR: No such account!")
                    return redirect(url_for('login'))

        # If the organization field is NOT EMPTY.
        else:
            if organisation in folders:
                orgPath = os.path.join(BASE_DIR, organisation)
                existEmail = db_Login(email, orgPath)
                emailCheck = db_Query(email, orgPath, column='email')
                # db_Query(email, f"{BASE_DIR}/{organisation}/", column='email')
                passCheck = db_Query(email, orgPath, column='password')
                if passCheck is not None:
                    print(f'The password match is {check_password_hash(passCheck, password)}')
                    data = json.load(open(f"{BASE_DIR}/{organisation}/credentials.json"))

            else:
                print('Organization does not exist.')
                return redirect(url_for('login'))

        key = E.decryptfile(cert, f"{BASE_DIR}/{organisation}/keys.txt")
        access_control = E.decryptaesfile(key, f"{BASE_DIR}/{organisation}/access.json")
        print(f'{email==emailCheck} {check_password_hash(passCheck, password)}')
        # if email in list(data.keys()) and password == data[email]["password"] and check_password_hash(passCheck, password):

        if email == emailCheck and check_password_hash(passCheck, password):
            print("SUCCESS LOGIN")

            # Generates a sessionID after login. (THIS SHOULD NOT HAPPEN NOW.)
            sessionID = generateSessionID()
            # session['sessionID'] = sessionID


            cert.stream.seek(0)
            certdict[sessionID] = cert.stream.read()
            session["sid"] = sessionID
            print(certdict, "JDPJSPDJSADPKJDAOODKOPASKDJPASJDPAJDOAPDJp")
            session["root"] = organisation
            session['filepage'] = f"{BASE_DIR}/{session['root']}"
            sessionDict[sessionID] = LogUser(email, password)
            # sessionDict[sessionID] = OrgUser("oooo",email, password,"klklkl")
            # resp = make_response(redirect(url_for('home')))
            resp = make_response(redirect(url_for('twoFA')))
            session['LoginSuccess'] = ''
            # session['AuthMethods'] = ''
            session['emailValid'] = email
            session['ac'] = access_control
            # resp.set_cookie('sessionID', sessionID)
            if access_control['type'] == "RBAC":
                role = db_Query(email, orgPath, column='role')
                resp.set_cookie('role', role)
                resp.set_cookie('email', email)
            elif access_control['type'] == "DAC":
                resp.set_cookie('email', email)
            return resp

    return render_template('login.html', reqorg=reqorg)


@app.route('/2fa', methods=['GET', 'POST'])
def twoFA():
    # Checks for valid login and AuthMethods page permission.
    # if 'LoginSuccess' not in session or 'AuthMethods' not in session:
    #     return render_template('error404.html')
    if 'LoginSuccess' not in session:
        return render_template('error404.html')
    # elif 'sessionID' in session:
    #     return redirect(url_for('home'))

    if request.method == 'POST':
        postData = request.form

        # After user picks his choice of 2FA.
        if 'vCodeType' in postData and postData['vCodeType'] != "":
            email = session['emailValid']

            # Passes valid authentication.
            session['vAuthentication'] = 'passed'

            # Using the user's email
            if postData['vCodeType'] == 'email':
                session['vCodeType'] = 'email'
                vCodeType = 'email'

                # Generates a code for Email 2FA.
                vCodeGenerated = generateVerificationCode()
                print('vCode: {}'.format(vCodeGenerated))
                emailSubject = 'Shake Shack - Timed OTP'
                user_msg = f'Do not share your OTP with anyone else. The code is {vCodeGenerated}.'

                if mailfea == "yes":
                    t1 = mailing(postData['email'], emailSubject, vCodeGenerated, user_msg)
                    t1.start()

                session['vCodeGenerated'] = vCodeGenerated

                return render_template('authMethods.html', session=session, vCodeType=vCodeType)

            elif postData['vCodeType'] == 'qrCode':
                session['vCodeType'] = 'qrCode'
                return redirect(url_for('login_2fa_form'))

            # if postData['resend'] == 'yes':
            #     session['resend'] = 'yes'
            #     session['vAuthentication'] = 'passed'
            # elif postData['resend'] == 'no':
            #     session['resend'] = 'no'

            elif postData['vCodeType'] == 'TOTP':
                session['vCodeType'] = 'TOTP'
                return redirect(url_for('login_2fa_form'))

            return render_template('abandon.html', session=session)

        # If code is entered
        elif 'vCode' in postData:
            # If email code EQUALS session code
            if session['vCodeGenerated'] == postData['vCode']:
                session['sessionID'] = session['sid']
                # print('\nONLINE NARABA\n')
                resp = redirect(url_for('home'))
                resp.set_cookie('sessionID', session["sessionID"])

                return resp

            else:
                session['vCodeAuthorization'] = 'failed'
                vCodeType = 'email'

                # Generates a code for Email 2FA, again.
                vCodeGenerated = generateVerificationCode()
                print('vCode: {}'.format(vCodeGenerated))
                emailSubject = 'Shake Shack - One-Time-Password'
                user_msg = f'Do not share your OTP with anyone else. The code is {vCodeGenerated}.'
                t1 = mailing(postData['email'], emailSubject, vCodeGenerated, user_msg)
                t1.start()

                session['vCodeGenerated'] = vCodeGenerated
                return render_template('authMethods.html', session=session, vCodeType=vCodeType)

        else:
            session['vAuthentication'] = 'Failed'

            return redirect(url_for('login'))

    else:
        return render_template('authMethods.html', session=session, vCodeType='')


@app.route('/2fa/totp/', methods=['GET', 'POST'])
def login_2fa_form():
    if 'sessionID' in session:
        return redirect(url_for('home'))

    if request.method == 'POST':
        # getting secret key used by user
        secret = request.form.get("secret")
        # getting OTP provided by user
        otp = int(request.form.get("otp"))

        print(f'Secret Key: {secret}\nOTP: {otp}')

        # verifying submitted OTP with PyOTP
        if pyotp.TOTP(secret).verify(otp):
            # inform users if OTP is valid
            print("The TOTP 2FA token is valid")
            session['sessionID'] = session["sid"]
            resp = redirect(url_for('home'))
            resp.set_cookie('sessionID', session["sessionID"])
            return resp
        else:
            # inform users if OTP is invalid
            print("You have supplied an invalid 2FA token!")
            return redirect(url_for("login_2fa_form"))

    if session['vCodeType'] == 'TOTP':
        secret = pyotp.random_base32()
        print(f'Secret: {secret}')
        print(f'\nThe TOTP is: {pyotp.TOTP(secret).now()}\n')

        emailSubject = 'Shake Shack - Secret Key (Google Auth)'
        user_msg = f'Do not share your OTP with anyone else. The code is {secret}.'
        t1 = mailing(session['emailValid'], emailSubject, secret, user_msg)
        t1.start()

        return render_template("totp_2fa.html", secret=secret, session=session)

    elif session['vCodeType'] == 'qrCode':
        print('New secret key has been printed.')
        secret = pyotp.random_base32()
        imgQR = qrcode.make(secret)
        imgQR_name = "qrCodeAuth.jpg"
        # imgQR.save(imgQR_name)
        imgQR.save(os.path.join(app.root_path, 'static/' + imgQR_name))
        print(f'\nThe OTP of this QR CODE is: {pyotp.TOTP(secret).now()}\n')

        # return render_template('totp_2fa.html', session=session)
        return render_template("totp_2fa.html", secret=secret, session=session)


@app.route('/verifyEmail', methods=['GET', 'POST'])
def verifyEmail():
    if 'RegisteredEmail' in session:
        email = session['RegisteredEmail']
    else:
        return render_template('error404.html')

    # Checks if the email is valid.
    email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    if re.fullmatch(email_regex, email):
        print('Valid Email.')

        orgPath = os.path.join(BASE_DIR, session['OrganisationName'])
        print(orgPath)
        existEmail = db_Query(email, orgPath, column='email')
        print(f'existEmail: {existEmail}')

        if existEmail:
            token = generate_confirmation_token(email)
            confirm_url = url_for('confirm_email', token=token, _external=True)
            emailSubject = 'Shake Shack - Email Verification'
            mailMSG = 'Thank you for signing up! we just need you to click the link to verify your account.\n\n' \
                      f'{confirm_url}\nIf this action was not your doing, please secure your account and ' \
                      f'lock it immediately. '
            t1 = mailing(email, emailSubject, False, mailMSG)
            t1.start()
            session['emailSent'] = 'A password reset will be sent to your email.'
            print('Email has been sent.')

    return render_template('verifyEmail.html', session=session)


@app.route('/verifyEmail/<token>', methods=['GET', 'POST'])
def confirm_email(token):
    email = confirm_token(token)

    if email is not False:
        print('SUCCESS: Confirmation link accessed!')
    else:
        print('FAILURE: The confirmation link is invalid or has expired.')
        return render_template('error404.html')

    print(f'The email is {email}.')
    # Gets the organization folder directory.
    drive_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drive')

    # Goes through folders, and checks for verification mark.
    folders = os.listdir(BASE_DIR)
    orgPath = os.path.join(BASE_DIR, session['OrganisationName'])
    verifyValue = db_Query(email, orgPath, column='verified')
    print(f'verifyValue: {verifyValue}')

    if verifyValue != 'False':
        print('Your account is verified.')

    else:
        verifyValue = 'True'
        db_Update(email, orgPath, column='verified', value=verifyValue)
        session['vSuccess'] = 'Verification Successful!'

    # for index in folders:
    #     # Makes a path to create the organization folder
    #     driveOrg_path = os.path.join(drive_path, index, 'orgDB.db')
    #
    #     verifyValue = db_Query(email, driveOrg_path, column='verified')
    #     if verifyValue != 'False':
    #         print('Your account is verified.')
    #
    #     else:
    #         verifyValue = 'True'
    #         db_Update(email, driveOrg_path, column='verified', value=verifyValue)
    #         session['vSuccess'] = 'Verification Successful!'

    # Clears the session after verification changes.
    session.clear()
    session['RegisteredUser'] = 'New user registered!'
    return redirect(url_for('login'))


@app.route('/passwordReset', methods=['GET', 'POST'])
def passwordReset():
    session.clear()
    existEmail = None
    postData = request.form
    if request.method == 'POST':
        email = postData['email']

        email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if re.fullmatch(email_regex, email):
            print('Valid Email.')

            # Gets the organization folder directory.
            folders = os.listdir(BASE_DIR)

            # Searches the folders for the DB you are in.
            for orgFolder in folders:
                # Makes a path to create the organization folder
                driveOrg_path = os.path.join(BASE_DIR, orgFolder)
                existEmail = db_Query(email, driveOrg_path, column='email')
                verifyCheck = db_Query(email, driveOrg_path, column='verified')
                if existEmail is None:
                    continue

                # Email is found in database.
                elif existEmail and verifyCheck == 'True':
                    token = generate_confirmation_token(email)
                    confirm_url = url_for('confirm_password', token=token, _external=True)
                    emailSubject = 'Shake Shack - Reset your password'
                    mailMSG = 'Here is the link for you to reset your password.\n\n' \
                              f'{confirm_url}\nIf this action was not your doing, please secure your account and ' \
                              f'lock it immediately. '
                    t1 = mailing(email, emailSubject, False, mailMSG)
                    t1.start()
                    session['emailSent'] = 'A password reset will be sent to your email.'
                    print('Email has been sent.')
                    return redirect(url_for('login'))

                elif existEmail and verifyCheck == 'False':
                    print('\nIt seems that your account is not verified. Please verify it before you can change your password.\n')
                    session['NotVerified'] = 'It seems that your account is not verified. Please verify it before you can change your password.'
                    session['RegisteredEmail'] = email
                    session['OrganisationName'] = orgFolder
                    return redirect(url_for('verifyEmail'))

            if existEmail is None:
                existEmail = False
                return render_template('passwordReset.html', existEmail=existEmail)


    return render_template('passwordReset.html', session=session, existEmail=existEmail)


@app.route('/passwordReset/<token>', methods=['GET', 'POST'])
def confirm_password(token):
    session.clear()
    if request.method == 'POST':
        try:
            email = confirm_token(token)
            print('SUCCESS: Confirmation link accessed!')
        except:
            print('FAILURE: The confirmation link is invalid or has expired.')

        # Takes in new and confirmed password fields.
        postData = request.form
        newPassword = postData['newPassword']
        conPassword = postData['conPassword']

        # Series of checks for invalid inputs.
        if newPassword != conPassword:
            return render_template('passwordChangeForm.html', session=session)

        # If any empty fields.
        elif newPassword == '' or conPassword == '':
            session['emptyField'] = 'The fields are empty.'
            return render_template('passwordChangeForm.html', session=session)

        # If newPassword == conPassword. ALL CORRECT.
        else:
            # Gets the organization folder directory.
            folders = os.listdir(BASE_DIR)

            for orgFolder in folders:
                # Makes a path to create the organization folder
                driveOrg_path = os.path.join(BASE_DIR, orgFolder)
                emailExists = db_Query(email, driveOrg_path, column='email')

                data = {email: {"password": newPassword, "role": "owner"}}

                credJSON = open(f"{driveOrg_path}/credentials.json", 'r')
                credJSON_Object = json.load(credJSON)
                credJSON.close()
                print(f'\n{credJSON_Object}\n')

                if emailExists:
                    passValue = db_Query(email, driveOrg_path, column='password')
                    print(f'The new password is: {newPassword}.')

                    # Changes password in SQLITE Database.
                    changePassword = generate_password_hash(newPassword)
                    db_Update(email, driveOrg_path, column='password', value=changePassword)

                    print(credJSON_Object[email]['password'])
                    credJSON_Object[email]['password'] = newPassword

                    credJSON = open(f"{driveOrg_path}/credentials.json", 'w')
                    json.dump(credJSON_Object, credJSON)
                    credJSON.close()

                    session['PassResetComplete'] = 'Your password has been changed.'

            return redirect(url_for('login'))

    return render_template('passwordChangeForm.html', session=session)


@app.route('/home/', methods=['POST', 'GET'])
def home():
    global uploadpath, url,upload
    #variables
    delete = False
    edit = False
    dir_delete = False
    root = False
    change_perm=False
    upload = False
    creation=False
    blacklist = ['access.json', 'credentials.json', 'orgDB.db', 'keys.txt', 'cert.crt']
    print(session['root'], "KKKKKKKKKKKKKKKKKKKKKKKK")
    session['modify'] = []
    print(request.cookies.get('role'), "UYUUUUUUUUUUUUUUUUUUUUUU")
    print(session["root"], "GGG")

    if session["root"] == None:
        if request.method == "POST":
            foldersname = []

            session["reqorg"] = request.form["org"]

            for element in os.listdir(BASE_DIR):
                elements = element.split(seperator)
                foldersname.append("#".join(elements))
            if session["reqorg"] in foldersname:
                print("YESSS")
                return redirect(url_for("login"))
            print(foldersname,session["reqorg"])
        return render_template('home.html')
    # path = request.path
    # # Joining the base and the requested path
    # abs_path = BASE_DIR + "/"+ req_path
    # print(abs_path)
    # print(path)
    else:
        print(session["filepage"],"FFFFFFFFFFFFFFF")
        if request.method == "POST":
            if request.form["opt"] == "filepath":
                filepage = request.form["filepage"]
                session["filepage"] += "/" + filepage
                print(session["filepage"],"Fsssssssssssssssss")
            elif request.form["opt"] == "back":
                print(session['filepage'].split("/"))
                filepage = session['filepage'].split("/")
                filepage.pop()
                print(filepage,"FGAdadad")
                session['filepage']="/".join(filepage)
                print(session['filepage'])
        # Check if path is a file and serve
        # if os.path.isfile(abs_path):
        #     print("SADSD")
        #     return send_from_directory(uploadpath, os.path.basename(abs_path),as_attachment=True)
        back = True
        print(len(session['filepage']))
        rootpage = f"{BASE_DIR}/{session['root']}"
        filepage = session['filepage']
        if len(filepage) <= len(rootpage):
            back = False
        # Show directory contents
        files =[]
        directory =[]
        i = 0
        print(os.listdir(session['filepage']))
        data = requestac()
        filepage =session["filepage"]
        print(session["filepage"],"JFGfff")
        filepath = filepage.replace(BASE_DIR + "/", "")
        print(filepath)
        for item in os.listdir(session['filepage']):

            if data["type"] == "RBAC":
                if request.cookies.get("email")== list(data["fidir"][(filepath+"/" + item)]["user"].keys())[0]:
                    permission = data["fidir"][(filepath+"/" + item)]["user"][request.cookies.get("email")]
                elif request.cookies.get('role') in list(data["fidir"][(filepath+"/" + item)]["roles"].keys()):

                    permission = data["fidir"][(filepath+"/" + item)]["roles"][ request.cookies.get('role')]
                else:
                    permission =data["fidir"][(filepath+"/" + item)]["others"]





                print(permission)
            elif data["type"] == "DAC":
                if request.cookies.get("email") == list(data["fidir"][(filepath + "/" + item)]["user"].keys())[0]:
                    permission = data["fidir"][(filepath + "/" + item)]["user"][request.cookies.get("email")]
                elif request.cookies.get('email') in list(data["fidir"][(filepath + "/" + item)]["roles"].keys()):

                    permission = data["fidir"][(filepath + "/" + item)]["roles"][request.cookies.get('email')]
                else:
                    permission = data["fidir"][(filepath + "/" + item)]["others"]




            if os.path.isdir(session['filepage'] +"/" + item):

                print(request.cookies.get('email'),list(data["fidir"][filepath]["roles"].keys()))
                if "R" in permission or "O" in permission:
                    if item not in blacklist:
                        directory.append(item)
                if "P" in permission or "O" in permission:
                    dir_delete = True

            if os.path.isfile(session['filepage'] +"/" + item):
                if "R" in permission or "O" in permission:
                    if item not in blacklist:
                        files.append([item, os.path.getsize(session['filepage'] +"/" + item),time.ctime(os.path.getmtime(session['filepage'] +"/" + item)) ,time.ctime(os.path.getctime(session['filepage'] +"/" + item)), i])
                if "F" in permission:
                    edit = True
                if "M" in permission:

                    session["modify"].append(item)
                    print(session["modify"])
                if "O" in permission or "D" in permission:
                    delete =True
        if data["type"] == "RBAC":
            if request.cookies.get("email") == list(data["fidir"][filepath]["user"].keys())[0]:
                permission_dir = data["fidir"][filepath]["user"][request.cookies.get("email")]
            elif request.cookies.get('role') in list(data["fidir"][filepath]["roles"].keys()):

                permission_dir = data["fidir"][filepath]["roles"][request.cookies.get('role')]
            else:
                permission_dir = data["fidir"][filepath]["others"]
        elif data["type"] == "DAC":
            if request.cookies.get("email") == list(data["fidir"][filepath]["user"].keys())[0]:
                permission_dir = data["fidir"][filepath]["user"][request.cookies.get("email")]
            elif request.cookies.get('email') in list(data["fidir"][filepath]["roles"].keys()):

                permission_dir = data["fidir"][filepath]["roles"][request.cookies.get('email')]
            else:
                permission_dir = data["fidir"][filepath]["others"]
        print("PERMDIR:", permission_dir)
        print(filepath)

        if "O" in permission_dir:
            change_perm = True
        if "F" in permission_dir:
            upload = True
        if "D" in permission_dir:
            creation = True

            i+=1
        print(upload)
        uploadpath = session['filepage']
        url = request.url
        print(files)
        if "modifypermission" in session:
            no = session["modifypermission"]
        else:
            no = "yes"
            session["modifypermission"] = ''

        if request.cookies.get("roles") == "root" or request.cookies.get("email") in session["root"]:
            root = True

        return render_template('home.html', files=files, directory =directory, back =back, delete =delete, edit= edit, no=no,
                               dir_delete=dir_delete,change_perm=change_perm,upload=upload,creation=creation, root = root)


@app.route("/removefile",methods=["POST","GET"])
def removefile():
    ac = requestac()
    file = f"{session['filepage']}/{request.args.get('thing')}"

    filepath = session["filepage"] + "/" + request.args.get('thing')

    filepath = filepath.replace(BASE_DIR + "/", "")

    if os.path.isfile(file):
        os.remove(file)
        del ac["fidir"][filepath]

    if os.path.isdir(file):

        for item in os.listdir(file):
            del ac["fidir"][filepath+"/"+item]
        del ac["fidir"][filepath]
        shutil.rmtree(file)
        certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
        acupdate(ac, f"{BASE_DIR}/{session['root']}/access.json", certpath)
        session['ac'] = requestac()
    return redirect(url_for("home"))


@app.route('/homepage', methods=['GET', 'POST'])
def homepage():
    global session
    print('SUCCESS: Redirecting to homepage...')
    # if 'sessionID' not in session and 'user_login' not in session:
    #     return render_template('error404.html')

    if request.method == 'POST':
        postData = request.form
        if 'logoutCall' in postData and postData['logoutCall'] == 'logout':
            session.clear()
            resp = make_response(redirect(url_for('login')))
            return resp

    return render_template('homepage.html')


@app.route('/logout', methods=['POST', 'GET'])
def logout():
    session.pop('sessionID')
    resp = redirect(url_for('login'))
    resp.set_cookie('sessionID', '', expires=0)
    return resp

@app.route("/upload",methods=["POST","GET"])
def upload():
    global upload
    ac = requestac()
    if request.method == 'POST':
        print("WOMENBUYIYASN")
        file = request.files['file']
        fname, file_extension = os.path.splitext(file.filename)
        print(file_extension)
        filename = file.filename
        # filename = secure_filename(file.filename)
        print(file.filename,"<><><<><><")
        print(f"{session['root']}/{filename} fsddddddddddd")
        if upload ==True:



            print(file.filename,"IIIUUUUUUUUUUUUUUU")
            if file.filename in os.listdir(session['filepage']) and file.filename in session['modify']:
                if os.path.exists(f"{session['filepage']}/{filename}"):
                    print("YYYYYYYYYYYYYYYYU")
                print(file.read(),"DDDDDD")
                print(f"{session['filepage']}/{filename}")

                with open(f"{session['filepage']}/{filename}", "wb") as existingfile:
                    file.seek(0)
                    existingfile.write(file.read())
                    session["modifypermission"] = 'yes'


            else:

                session["modifypermission"] = 'yes'
                filepath = session["filepage"] + "/" + file.filename
                filepath = filepath.replace(BASE_DIR + "/", "")
                print(filepath,'filepath')
                if ac["type"] == "RBAC":
                    ac["fidir"].update({filepath: {"user": {request.cookies.get("email"): "OFMRD"},"roles": {"root": "OFMRD"}, "others": "-----"}})
                elif ac["type"] == "DAC":
                    ac["fidir"].update({filepath: {"user": {request.cookies.get("email"): "OFMRD"},"users": {session['root'].split(seperator)[0]: "OFMRD"},"others": "-----"}})
                certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
                acupdate(ac, f"{BASE_DIR}/{session['root']}/access.json", certpath)
                print("YESSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSSS")
                session['ac'] = requestac()
                file.save(f"{session['filepage']}/{filename}")
                if file_extension.lower() in [".png", ".jpg", ".jpeg"]:
                    print("SUCESS")
                    startwatermark(f"{session['filepage']}/{filename}", session["root"])

                print(decodewatermark(f"{session['filepage']}/{filename}", len(session["root"]*8)))

        else:
            session["modifypermission"] = 'no'
    return redirect(url_for("home"))


@app.route('/homepage/backupPage', methods=['GET', 'POST'])
def backupPage():
    global session
    postData = request.form
    print('Directing to Backup Page...')
    if request.method == 'POST':
        if postData['backupMode'] == 'dropboxUp':
            dbBackupPages()
            print('Pages Backed UP')
            return redirect(url_for('homepage'))


    return render_template('backupDB.html', session=session)


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


@app.route('/download', methods=['GET', 'POST'])
def download():
    print('/DOWNLOAD', request.full_path)
    filename = request.args.get('file')
    if filename.endswith('.docx') and 'mask' in request.args:
        doc = Document(session['filepage'] + '/' + filename)

        if not request.args.get('email') == 'on':
            email = re.compile('''(?:[a-z0-9!#$%&'*+/=?^_`{|}~-]+(?:\.[a-z0-9!#$%&'*+/=?^_`{|}~-]+)*|"(?:[\x01-\x08\x0b\x0c\x0e-\x1f\x21\x23-\x5b\x5d-\x7f]|\\[\x01-\x09\x0b\x0c\x0e-\x7f])*")@(?:(?:[a-z0-9](?:[a-z0-9-]*[a-z0-9])?\.)+[a-z0-9](?:[a-z0-9-]*[a-z0-9])?|\[(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?|[a-z0-9-]*[a-z0-9]:(?:[\x01-\x08\x0b\x0c\x0e-\x1f\x21-\x5a\x53-\x7f]|\\[\x01-\x09\x0b\x0c\x0e-\x7f])+)\])''')
            docx_replace_regex(doc, email, '<EMAIL_REDACTED>')

        if not request.args.get('phone') == 'on':
            phone = re.compile('[0-9]{8}')
            docx_replace_regex(doc, phone, '<PHONE_REDACTED>')

        if not request.args.get('nric') == 'on':
            nric = re.compile('^[STFG]\d{7}[A-Z]$')
            docx_replace_regex(doc, nric, '<NRIC_REDACTED>')

        fileStream = io.BytesIO()
        doc.save(fileStream)
        fileStream.seek(0)

        return send_file(fileStream, as_attachment=True, attachment_filename=filename)




    # abs_path = BASE_DIR+"/"+ session['filepage']
    # print(abs_path, session['filepage'] , "POPPPPPPPPPPPPPPPPPPPPPPPPPPP")
    uploads = f"{current_app.root_path}/{app.config['UPLOAD_FOLDER']}"
    return send_from_directory(session['filepage'], filename, as_attachment=True)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if 'sessionID' in session:
        return redirect(url_for('home'))

    result = ""
    if request.method == 'POST':
        userfold = "./ISPJ uploads/user"
        email = request.form["Email"]
        password = request.form["Password"]
        name = request.form["Profile Name"]
        phonenumber = request.form["Phone Number"]
        org = request.form["Organisation"]
        ac = request.form["ac"]
        # acp = request.form["mac"]
        email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if not re.fullmatch(email_regex, email):
            return render_template('register.html',  passwordok='Invalid email')

        # If the password does not pass the password policy.
        if password_policy(password)['password_ok'] == False:
            return render_template('register.html',  passwordok='notok')

        print(email, password, name, phonenumber, org)

        # List of emails
        emails = []
        user = OrgUser(email, name, password, phonenumber, False)

        if "!" in email:
            result = "Unacceptable characters : '!'"
            return render_template('register.html', result=result)

        # Retrieves the emails in the organization folder names.
        for dir in os.listdir(userfold):
            emails.append(dir.split(seperator)[0])

        # Checks if the organization email exists.
        if email.lower() in emails:
            result = "Account already exist!"
        else:
            if org != "":
                foldername = email.lower() +seperator+org.lower()
                path = userfold+"/"+foldername
                os.mkdir(path)
                print("Directory '% s' created" % org)
                # with open(f"{path}/credentials.txt", 'w') as f:
                #     f.write(f'{email}{seperator}{password}{seperator}owner\n')


            else:
                foldername = email.lower() +seperator+org.lower()
                path = userfold+"/"+foldername
                os.mkdir(path)
                print("Directory '% s' created" % email)
                # with open(f"{path}/credentials.txt", 'w') as f:
                #     f.write(f'{email}{seperator}{password}')

            db_Create(path)
            user.set_role('root')
            db_Register(user, path)
            if org != '':
                print("Directory '% s' created" % org)
            else:
                print(f'Private directory created for {name}')



            session['RegisteredEmail'] = email
            session['OrganisationName'] = foldername

            data = {email: {"password": password, "role": "root"}}


            open(f"{path}/credentials.json", 'x')
            json.dump(data, open(f"{path}/credentials.json", 'w+'))
            print('############')

            if ac == "RBAC":
                dataA = {"type":ac ,"fidir": {foldername:{"user":{email:"OFDRP"}, "roles": {"root":"OFDRP"}, "others": "---R-" }}, "classifier":["root"]}
            elif ac == "DAC":
                # dataA = {"type":ac ,email: {foldername: "OFMRD" }}
                dataA = {"type": ac,"fidir": {foldername: {"user": {email: "OFDRP"}, "roles": {email: "OFDRP"}, "others": "---R-"}},"classifier":[email]}
            # if acp == "YES":
            #     dataB = {"HS":[foldername], "S":[foldername], "I":[foldername], "P":[foldername]}
            #     json.dump(dataB, open(f"{path}/mac.json", 'w+'))
            open(f"{path}/access.json", 'x')
            key = E.generateaeskey()

            json.dump(dataA, open(f"{path}/access.json", 'w+'))
            print('############')
            accessdata = json.load(open(f"{path}/access.json"))

            for item in os.listdir(path):
                addfile(accessdata,foldername+"/"+item,email,"OFMRD", "root", "OFMRD","-----", "no",f"{path}/access.json")
            addfile(accessdata, foldername + "/" + "keys.txt", email, "OFMRD", "root", "OFMRD", "-----"
                    , "no",f"{path}/access.json")
            addfile(accessdata, foldername + "/" + "cert.crt", email, "OFMRD", "root", "OFMRD", "-----"
                    , "no", f"{path}/access.json")

            E.cert_gen_user(email,email,password, path,key)
            ##### FOR MAILING FEATURE DO NOT DELETE ######
            ##### FOR MAILING FEATURE DO NOT DELETE ######
            mkc.send_mail("shakeshackproject@gmail.com","Syocyno@gmail.com","KEYS AND CERT FOR SIMPLEBOARD", email)
            ##### FOR MAILING FEATURE DO NOT DELETE ######
            ##### FOR MAILING FEATURE DO NOT DELETE ######
            # with open(f"{path}/access.txt", 'w') as a:
            #     a.write(f"o----{seperator}{email}")
            print('Registration successful. Redirecting to email verification.')
            return redirect(url_for("verifyEmail"))

    return render_template('register.html', result = result)

@app.route('/regPasswordPolling', methods=['POST'])
def regPasswordPolling():
    postData = request.form['password']
    policyresult = password_policy(postData)
    print(postData)
    return jsonify('', render_template('registerPolling.html', policyresult=policyresult))


@app.route('/createfolder', methods=['GET', 'POST'])
def createfolder():
    global url
    if request.method == 'POST':
        folder = request.form["Folder"]
        path = f"{session['filepage']}/{folder}"
        filepath = session["filepage"] + "/" + folder
        filepath = filepath.replace(BASE_DIR + "/", "")
        os.mkdir(path)
        ac = requestac()
        print(ac)
        if ac["type"] == "RBAC":
            ac["fidir"].update({filepath:{"user":{request.cookies.get("email"):"OFDRP"}, "roles": {"root":"OFDRP"}, "others": "----R" }})
        elif ac["type"] == "DAC":
            ac["fidir"].update({filepath:{"user":{request.cookies.get("email"):"OFDRP"}, "roles": {session['root'].split(seperator)[0]:"OFDRP"}, "others": "----R" }})

        print(ac)
        certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
        acupdate(ac, f"{BASE_DIR}/{session['root']}/access.json", certpath)
        session['ac'] = requestac()
        print(session['ac'])

    return redirect(url_for("home"))

@app.route('/dirpermission', methods=['POST', 'GET'])
def dirpermission():
    session['ac'] = requestac()
    data = session['ac']
    perms = []
    filepath = session["filepage"] + "/" + request.args.get("thing")
    filepath = filepath.replace(BASE_DIR + "/", "")


    if data["type"] == "RBAC":
        classifier = data["classifier"]

    elif data["type"] == "DAC":
        classifier = data["classifier"]
    if request.method == 'POST':

        exist = False
        if data["type"] == "RBAC":
            role = request.form['role']

        elif data["type"] == "DAC":
            role = request.form['role']
        option = request.form['option']
        print(session["filepage"], request.args.get("thing"))
        filepath = session["filepage"] + "/" + request.args.get("thing")
        filepath = filepath.replace(BASE_DIR + "/", "")

        if "afile" in request.form:
            afile = request.form['afile']
        else:
            afile = "-"
        if "adir" in request.form:
            adir = request.form['adir']
        else:
            adir = "-"
        if "Read" in request.form:
            read = request.form['Read']
        else:
            read = "-"
        if "Delete" in request.form:
            delete = request.form['Delete']
        else:
            delete = "-"

        if option == "permission":
            data = session['ac']
            data["fidir"][filepath]["roles"][role] = f"-{afile}{adir}{read}{delete}"

            path = f"{BASE_DIR}/{session['root']}/access.json"
            certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
            acupdate(data, path, certpath)
            session['ac'] = requestac()
            # with open(f"{BASE_DIR}/{filepath.split('/')[3]}/access.txt", 'r+') as access:
            #     accesscontent = access.readlines()
            #     for i, line in enumerate(accesscontent):
            #         if filepath in line:
            #             accesscontent[i] += f"{permseperator}-{fc}{modify}{read}{delete}{seperator}{role}"
            #             exist = True
            #
            #     if exist != True:
            #         access.write(f"{filepath}: -{fc}{modify}{read}{delete}{seperator}{role}\n")
            #     else:
            #         access.seek(0)
            #         for line in accesscontent:
            #             access.write(line)
            print(data, "THE DAT CHACK")
    owner = list(data["fidir"][filepath]["user"].keys())[0]
    permrole = list(data["fidir"][filepath]["roles"].keys())
    permrole.insert(0, owner)
    perms.append(list(data["fidir"][filepath]["user"].values())[0])
    for values in list(data["fidir"][filepath]["roles"].values()):
        perms.append(values)
    perms.append(data["fidir"][filepath]["others"])
    return render_template("dirpermission.html", ac=data["type"], roles=classifier, owner=owner, permroles = permrole, perms=perms)

@app.route('/permission', methods=['GET', 'POST'])
def permission():
    global uploadpath


    session['ac'] = requestac()
    data = session['ac']
    perms = []
    filepath = session["filepage"] + "/" + request.args.get("thing")
    filepath = filepath.replace(BASE_DIR + "/", "")
    owner = list(data["fidir"][filepath]["user"].keys())[0]
    permrole =  list(data["fidir"][filepath]["roles"].keys())
    permrole.insert( 0,owner)
    perms.append(list(data["fidir"][filepath]["user"].values())[0])
    for values in list(data["fidir"][filepath]["roles"].values()):
        perms.append(values)
    perms.append(data["fidir"][filepath]["others"])
    if data["type"] == "RBAC":
        classifier = data["classifier"]

    elif data["type"] == "DAC":
        classifier = data["classifier"]
    if request.method == 'POST':

        exist = False
        if data["type"] == "RBAC":
            role = request.form['role']

        elif data["type"] == "DAC":
            role = request.form['email']
        option = request.form['option']
        print(session["filepage"],request.args.get("thing"))
        filepath = session["filepage"] +"/" + request.args.get("thing")
        filepath = filepath.replace(BASE_DIR +"/","")
        if "Full Control" in request.form:
            fc = request.form['Full Control']
        else:
            fc = "-"
        if "Modify" in request.form:
            modify = request.form['Modify']
        else:
            modify = "-"
        if "Read" in request.form:
            read = request.form['Read']
        else:
            read = "-"
        if "Delete" in request.form:
            delete = request.form['Delete']
        else:
            delete = "-"

        if option == "permission":
            data = session['ac']
            data["fidir"][filepath]["roles"][role] = f"-{fc}{modify}{read}{delete}"

            path = f"{BASE_DIR}/{session['root']}/access.json"
            certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
            acupdate(data, path, certpath)
            session['ac'] = requestac()
                    # with open(f"{BASE_DIR}/{filepath.split('/')[3]}/access.txt", 'r+') as access:
            #     accesscontent = access.readlines()
            #     for i, line in enumerate(accesscontent):
            #         if filepath in line:
            #             accesscontent[i] += f"{permseperator}-{fc}{modify}{read}{delete}{seperator}{role}"
            #             exist = True
            #
            #     if exist != True:
            #         access.write(f"{filepath}: -{fc}{modify}{read}{delete}{seperator}{role}\n")
            #     else:
            #         access.seek(0)
            #         for line in accesscontent:
            #             access.write(line)
            print(data,"THE DAT CHACK")
    return render_template("permission.html", ac=data["type"], roles=classifier, owner=owner, permroles = permrole, perms=perms)

@app.route('/assignrole', methods=['GET', 'POST'])
def assignrole():
    global certdict

    session['ac'] = requestac()
    data = session['ac']


    if data["type"] == "RBAC":
        classifier = data["classifier"]

        if request.method == "POST":
            opt = request.form["opt"]
            if opt == "assign":
                # dataC = json.load(open(f"{BASE_DIR}/{session['root']}/credentials.json"))
                #
                email = request.form["email"]
                password = request.form["password"]
                role = request.form["role"]
                # dataC[email] =  {"password": password, "role": role}
                user = OrgUser(email, email.split("@")[0], password, "31223214", True)
                user.set_role(role)
                print(f"{BASE_DIR}/{session['root']}/orgDB.db","PATH LAH CHIBAI")
                db_Register(user,f"{BASE_DIR}/{session['root']}" )
            if opt == "create":
                dataA = session['ac']
                role = request.form["role"]
                dataA["classifier"].append(role)
                path = f"{BASE_DIR}/{session['root']}/access.json"
                certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
                acupdate(dataA, f"{BASE_DIR}/{session['root']}/access.json",certpath)

                  # with open(f"{BASE_DIR}/{session['root']}/credential.txt", "r+"):
            #     print()
    elif data["type"] == "DAC":
        classifier = data["classifier"]
        if request.method == "POST":



            email = request.form["email"]
            password = request.form["password"]


            user = OrgUser(email, email.split("@")[0], password, "31223214", True)

            db_Register(user,f"{BASE_DIR}/{session['root']}" )
            dataA = session['ac']
            role = request.form["email"]
            dataA["classifier"].append(role)
            path = f"{BASE_DIR}/{session['root']}/access.json"
            certpath = f"{BASE_DIR}/{session['root']}/cert.crt"
            acupdate(dataA, f"{BASE_DIR}/{session['root']}/access.json", certpath)

    return render_template("assign.html", roles = classifier, ac = data["type"] )


class mailing(threading.Thread):

    def __init__(self, email, subject, code, user_msg):
        threading.Thread.__init__(self)
        self.__email = email
        self.__subject = subject
        self.__code = code
        self.__user_msg = user_msg

    def run(self):
        with app.app_context():
            msg = Message(self.__subject, sender='ShakeShackproject@gmail.com',
                          recipients=[self.__email])
            # msg.body = 'Your authentication code is {}'.format(self.__code)
            print(self.__subject, self.__email)
            msg.body = f'Hello {self.__email}!\n{self.__user_msg}'
            print(msg.body)
            global ENABLE_MAIL
            if ENABLE_MAIL:
                mail.send(msg)

def addfile(data ,foldername,email,permission1,role,permission2,permission3,exist, path):
    if exist == "no":


        data["fidir"].update({foldername:{"user":{email:permission1}, "roles": {role:permission2}, "others": permission3 }})
        json.dump(data, open(path,"w+"))
    elif exist == "yes":
        data["fidir"][foldername] ={"user":{email:permission1}, "roles": {role:permission2}, "others": permission3 }
# Generates Verification Code for OTP
def generateVerificationCode():
    code = ''
    for _ in range(6):
        code += str(randint(0, 9))

    return code

@app.route("/waterm",methods=["POST","GET"])
def waterm():
    validity = ""
    if request.method == "POST":
        file = request.files["water"]
        filename = file.filename
        file.save(f"./tempfiles/{filename}")

        validity=decodewatermark("C:/NYP stuff/ISPJ-Repo/ISPJ uploads/user/lel@gmail.com!$%()gg/catwide.jpg",len("lel@gmail.com!$%()gg")*8)

    return render_template("waterm.html", validity =validity)

@app.route('/ping', methods=['GET'])
def ping():
    memory = psutil.Process(os.getpid()).memory_info().rss / 1024 ** 2
    return json.dumps({}), 200, {'memory_usage': memory}

# @app.route("/waterm",methods=["POST","GET"])
# def waterm():
#     validity = ""
#     if request == "POST":
#         file = request.files["water"]
#
#         validity=decodewatermark(file,len(session["root"]))
#     return render_template("waterm.html", validity =validity)
# Request Access Control?
def requestac():
    print(certdict)
    cert = certdict[request.cookies.get("sessionID")]
    key = E.decryptfile(cert, f"{BASE_DIR}/{session['root']}/keys.txt",Yes="yes")
    access_control = E.decryptaesfile(key, f"{BASE_DIR}/{session['root']}/access.json")
    return access_control

def acupdate(plaintext, filepath,certpath):
    cert = certdict[request.cookies.get("sessionID")]
    key = E.decryptfile(cert, f"{BASE_DIR}/{session['root']}/keys.txt",Yes="yes")
    E.encryptaes(key, str(plaintext).encode("ISO-8859-1"),filepath)





if __name__ == '__main__':
    if VERIFY_FILES:
        export_backup()
        dbBackupPages()
        # backup_pages()
        # validate_pages(initialize=True)
        # dt.start()
    app.run(debug=True, port=5000)
